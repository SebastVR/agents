import os
import re
from typing import List

import boto3
import litellm
from config.config import settings
from crewai import Agent
from crewai.llm import LLM
from crewai.tools import tool
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor

# from langchain.chat_models import ChatOpenAI

# from langchain.chat_models import ChatOpenAI

# OPENAI_API_KEY = settings.openai_api_key
AWS_ACCESS_KEY_ID = settings.aws_access_key_id
AWS_SECRET_ACCESS_KEY = settings.aws_secret_access_key
AWS_DEFAULT_REGION = settings.aws_default_region
MODEL_NAME = settings.model_name
# litellm.set_verbose = True
# litellm.api_base = "https://bedrock-runtime.us-east-1.amazonaws.com"

client = boto3.client(
    "bedrock-runtime",
    region_name=AWS_DEFAULT_REGION,
    aws_access_key_id=AWS_ACCESS_KEY_ID,
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
)


@tool
def extraer_headings_desde_docx(input_dir: str) -> List[dict]:
    """
    Extrae los encabezados con su numeración desde documentos Word (.docx) en el directorio especificado.
    """
    headings = []

    print(f"📁 Ruta recibida: {input_dir}")
    print(f"📂 Archivos en el directorio:")
    if not os.path.exists(input_dir):
        print("❌ El directorio no existe.")
        return []
    print(os.listdir(input_dir))
    archivos = sorted(
        [
            f
            for f in os.listdir(input_dir)
            if f.endswith(".docx") and not f.startswith("~$")
        ]
    )
    for filename in archivos:
        path = os.path.join(input_dir, filename)
        print(f"\n📄 Procesando archivo: {filename}")
        doc = Document(path)
        for para in doc.paragraphs:
            estilo = para.style.name.lower()
            texto = para.text.strip()
            if "tabla de contenido" in texto.lower() or not texto:
                continue
            if estilo.startswith("heading"):
                try:
                    nivel = int(estilo.replace("heading", "").strip()) - 1
                except ValueError:
                    continue
                match = re.match(r"^((\d+\.?)+)\s+(.*)$", texto)
                if match:
                    numeracion = match.group(1)
                    titulo = match.group(3)
                    texto_final = f"{numeracion} {titulo}"
                    print(f"✅ Heading numerado: {texto_final} (Nivel {nivel})")
                else:
                    texto_final = texto
                    print(f"⚠️ Heading sin numeración: {texto} (Nivel {nivel})")
                headings.append({"text": texto_final, "level": nivel})

    print("\n📚 Encabezados extraídos:")
    for h in headings:
        print(f"- {h['text']} (Nivel {h['level']})")

    return headings


@tool
def crear_doc_tabla_contenido(headings: List[dict], output_file: str):
    """
    Genera un documento Word con la tabla de contenido a partir de una lista de encabezados.
    """
    print("➡️ Recibido en crear_doc_tabla_contenido:")
    print(f"📌 output_file = {output_file}")
    print(f"📌 Total headings = {len(headings)}")

    doc = Document()
    if "TOC_Entry" not in [s.name for s in doc.styles]:
        toc_style = doc.styles.add_style("TOC_Entry", WD_STYLE_TYPE.PARAGRAPH)
        toc_style.font.name = "Calibri"
        toc_style.font.size = Pt(11)
        toc_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_heading("Tabla de Contenido", level=0)

    for entry in headings:
        level = entry.get("level", 0)
        p = doc.add_paragraph(style="TOC_Entry")
        p.paragraph_format.left_indent = Pt(18 * level)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(entry["texto"] if "texto" in entry else entry["text"])

    doc.save(output_file)
    print(f"\n✅ Tabla de contenido generada en: {output_file}")


def crear_agente_toc():
    # Configurar litellm para usar Bedrock
    # litellm.set_verbose = True
    os.environ["LITELLM_LOG"] = "DEBUG"
    litellm.api_base = "https://bedrock-runtime.us-east-1.amazonaws.com"  # Asegúrate que coincida con tu región

    llm = LLM(
        model=settings.model_name,
        api_base=litellm.api_base,
        # provider="bedrock",
        aws_access_key_id=settings.aws_access_key_id,
        aws_secret_access_key=settings.aws_secret_access_key,
        aws_region_name=settings.aws_default_region,
        temperature=0.7,
        max_tokens=1000,
        drop_params=True,  # importante para Bedrock
    )

    # Prueba rápida del modelo
    try:
        respuesta = llm.call("¿Cuál es la capital de Colombia?")
        print("🧠 Respuesta del modelo Bedrock:", respuesta)
    except Exception as e:
        print("❌ Error al conectarse con Bedrock:", e)

    return Agent(
        role="Generador de Tabla de Contenido",
        goal="Leer documentos Word y construir tabla de contenido conservando numeración original",
        backstory=(
            "Eres un asistente experto en documentación técnica. Tu tarea es analizar documentos Word "
            "y generar una tabla de contenido clara y bien estructurada usando la numeración original de los encabezados."
        ),
        verbose=True,
        allow_delegation=False,
        tools=[extraer_headings_desde_docx, crear_doc_tabla_contenido],
        llm=llm,
    )


#####################################
#######################################
